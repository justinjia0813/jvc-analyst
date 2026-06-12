#!/usr/bin/env python3
from __future__ import annotations

import json
import subprocess
import sys
import tempfile
from zipfile import ZipFile
from xml.etree import ElementTree as ET
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
GENERATOR = ROOT / "skills/jvc-meeting-notes/scripts/generate_meeting_notes.py"
DEFAULT_TEMPLATE = ROOT / "skills/jvc-meeting-notes/templates/访谈纪要模板.docx"
EXPECTED_PAGE = (21.0, 29.7)
EXPECTED_MARGINS = (2.54, 2.54, 3.17, 3.17)
W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
EXPECTED_SECTION_XML = {
    "pgSz": {"w": "11906", "h": "16838"},
    "pgMar": {
        "top": "1440",
        "right": "1800",
        "bottom": "1440",
        "left": "1800",
        "header": "851",
        "footer": "992",
        "gutter": "0",
    },
    "cols": {"space": "425", "num": "1"},
    "docGrid": {"type": "lines", "linePitch": "312", "charSpace": "0"},
}
EXPECTED_ROLE_FORMATS = [
    ("title", "center", "36", True),
    ("section", "both", "20", True),
    ("body", "both", "20", False),
    ("section", "both", "20", True),
    ("subsection", "both", "20", False),
    ("body", "both", "20", False),
]


def write_json(path: Path, data: dict) -> None:
    path.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")


def run_generator(data: Path, output: Path) -> None:
    subprocess.run(
        [
            sys.executable,
            str(GENERATOR),
            str(data),
            "--template",
            str(DEFAULT_TEMPLATE),
            "--output",
            str(output),
        ],
        cwd=ROOT,
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )


def rounded_section(doc: Document) -> tuple[tuple[float, float], tuple[float, float, float, float]]:
    section = doc.sections[0]
    page = (round(section.page_width.cm, 2), round(section.page_height.cm, 2))
    margins = (
        round(section.top_margin.cm, 2),
        round(section.bottom_margin.cm, 2),
        round(section.left_margin.cm, 2),
        round(section.right_margin.cm, 2),
    )
    return page, margins


def nonempty_styles(path: Path) -> list[str]:
    return [p.style.name for p in Document(path).paragraphs if p.text.strip()]


def document_root(path: Path) -> ET.Element:
    with ZipFile(path) as archive:
        return ET.fromstring(archive.read("word/document.xml"))


def w_attr(element: ET.Element, name: str) -> str | None:
    return element.get(f"{W}{name}")


def assert_attrs(element: ET.Element | None, expected: dict[str, str], label: str) -> None:
    assert element is not None, f"missing {label}"
    actual = {key: w_attr(element, key) for key in expected}
    assert actual == expected, f"{label} mismatch: {actual} != {expected}"


def paragraph_text(paragraph: ET.Element) -> str:
    return "".join(node.text or "" for node in paragraph.findall(f".//{W}t"))


def nonempty_paragraph_xml(path: Path) -> list[ET.Element]:
    body = document_root(path).find(f"{W}body")
    assert body is not None, "missing document body"
    return [
        paragraph
        for paragraph in body.findall(f"{W}p")
        if paragraph_text(paragraph).strip()
    ]


def first_text_run(paragraph: ET.Element) -> ET.Element:
    for run in paragraph.findall(f"{W}r"):
        if "".join(node.text or "" for node in run.findall(f".//{W}t")).strip():
            return run
    raise AssertionError("paragraph has no text run")


def assert_standard_section_xml(path: Path) -> None:
    root = document_root(path)
    sect = root.find(f".//{W}sectPr")
    assert sect is not None, "missing section properties"
    assert_attrs(sect.find(f"{W}pgSz"), EXPECTED_SECTION_XML["pgSz"], "page size XML")
    assert_attrs(sect.find(f"{W}pgMar"), EXPECTED_SECTION_XML["pgMar"], "page margin XML")
    assert_attrs(sect.find(f"{W}cols"), EXPECTED_SECTION_XML["cols"], "columns XML")
    assert_attrs(sect.find(f"{W}docGrid"), EXPECTED_SECTION_XML["docGrid"], "docGrid XML")
    header_refs = sect.findall(f"{W}headerReference")
    assert len(header_refs) == 1, f"expected one empty default header reference, got {len(header_refs)}"
    assert w_attr(header_refs[0], "type") == "default", "header reference should be default"


def assert_role_format(paragraph: ET.Element, expected: tuple[str, str, str, bool], label: str) -> None:
    role, expected_jc, expected_size, expected_bold = expected
    p_pr = paragraph.find(f"{W}pPr")
    assert p_pr is not None, f"{label} {role} paragraph missing pPr"
    assert p_pr.find(f"{W}pStyle") is None, f"{label} {role} should use implicit Normal style"
    jc = p_pr.find(f"{W}jc")
    assert jc is not None and w_attr(jc, "val") == expected_jc, (
        f"{label} {role} alignment mismatch"
    )

    run = first_text_run(paragraph)
    r_pr = run.find(f"{W}rPr")
    assert r_pr is not None, f"{label} {role} run missing rPr"
    r_fonts = r_pr.find(f"{W}rFonts")
    assert r_fonts is not None, f"{label} {role} run missing rFonts"
    assert w_attr(r_fonts, "ascii") == "Times New Roman", f"{label} {role} ascii font mismatch"
    assert w_attr(r_fonts, "hAnsi") == "Times New Roman", f"{label} {role} hAnsi font mismatch"
    assert w_attr(r_fonts, "eastAsia") == "KaiTi", f"{label} {role} eastAsia font mismatch"
    sz = r_pr.find(f"{W}sz")
    assert sz is not None and w_attr(sz, "val") == expected_size, (
        f"{label} {role} font size mismatch"
    )
    bold = r_pr.find(f"{W}b")
    if expected_bold:
        assert bold is not None and w_attr(bold, "val") not in {"0", "false", "False"}, (
            f"{label} {role} should be bold"
        )
    else:
        assert bold is not None and w_attr(bold, "val") in {"0", "false", "False"}, (
            f"{label} {role} should be explicitly non-bold"
        )


def assert_standard_paragraph_xml(path: Path, label: str) -> None:
    paragraphs = nonempty_paragraph_xml(path)
    assert len(paragraphs) >= len(EXPECTED_ROLE_FORMATS), (
        f"{label} should have enough paragraphs for role checks"
    )
    for paragraph, expected in zip(paragraphs, EXPECTED_ROLE_FORMATS):
        assert_role_format(paragraph, expected, label)


def main() -> int:
    with tempfile.TemporaryDirectory() as tmpdir:
        tmp = Path(tmpdir)
        meeting_json = tmp / "meeting.json"
        talk_json = tmp / "talk.json"
        meeting_doc = tmp / "meeting.docx"
        talk_doc = tmp / "talk.docx"

        write_json(
            meeting_json,
            {
                "title": "2026/06/12 线上 访谈{格式测试}",
                "filename": "meeting.docx",
                "sections": [
                    {"heading": "一、公司基本情况", "content": "公司简介。"},
                    {
                        "heading": "二、公司核心技术",
                        "subsections": [
                            {"heading": "技术路线", "content": "技术内容。"},
                        ],
                    },
                ],
            },
        )
        write_json(
            talk_json,
            {
                "title": "2026/06/12 线上 客户访谈{格式测试}",
                "filename": "talk.docx",
                "sections": [
                    {"heading": "一、访谈基本信息", "content": "日期、形式、受访人。"},
                    {
                        "heading": "二、问答纪要",
                        "subsections": [
                            {"heading": "Q1：客户如何使用？", "content": "问题：...\n回答摘要：..."},
                        ],
                    },
                ],
            },
        )

        run_generator(meeting_json, meeting_doc)
        run_generator(talk_json, talk_doc)

        for label, path in {"meeting": meeting_doc, "talk": talk_doc}.items():
            page, margins = rounded_section(Document(path))
            assert page == EXPECTED_PAGE, f"{label} page size mismatch: {page}"
            assert margins == EXPECTED_MARGINS, f"{label} margin mismatch: {margins}"
            assert_standard_section_xml(path)
            assert_standard_paragraph_xml(path, label)
            styles = nonempty_styles(path)
            assert set(styles) == {"Normal"}, f"{label} should use only Normal paragraphs: {styles}"

        meeting_styles = nonempty_styles(meeting_doc)
        talk_styles = nonempty_styles(talk_doc)
        assert meeting_styles == talk_styles, (
            "meeting and talk should share the same Normal/direct-formatting structure: "
            f"{meeting_styles} != {talk_styles}"
        )

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
