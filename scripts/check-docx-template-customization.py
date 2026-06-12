#!/usr/bin/env python3
from __future__ import annotations

import json
import os
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
GENERATOR = ROOT / "skills/jvc-meeting-notes/scripts/generate_meeting_notes.py"
AUTO_TEMPLATE = ROOT / "skills/jvc-meeting-notes/templates/custom.docx"


def make_template(path: Path, marker: str) -> None:
    doc = Document()
    section = doc.sections[0]
    section.header.paragraphs[0].text = marker
    doc.add_paragraph("示例标题")
    doc.add_paragraph("一、示例章节")
    doc.add_paragraph("示例正文")
    doc.save(path)


def make_data(path: Path) -> None:
    data = {
        "title": "2026/06/12 线上 访谈{模板测试}",
        "filename": "模板测试.docx",
        "sections": [
            {"heading": "一、基本信息", "content": "这是正文。"},
            {
                "heading": "二、问答纪要",
                "subsections": [
                    {
                        "heading": "Q1：问题",
                        "content": "完整回答：测试。\n对应事实层维度：其他\n待验证点：无明显待验证点。",
                    },
                ],
            },
        ],
    }
    path.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")


def output_header_text(path: Path) -> str:
    doc = Document(path)
    parts: list[str] = []
    for section in doc.sections:
        for header in (section.header, section.first_page_header, section.even_page_header):
            parts.extend(paragraph.text for paragraph in header.paragraphs)
    return "\n".join(parts)


def run_generator(data: Path, output: Path, env: dict[str, str]) -> None:
    subprocess.run(
        [sys.executable, str(GENERATOR), str(data), "--output", str(output)],
        cwd=ROOT,
        env=env,
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )


def assert_uses_env_template(tmp: Path) -> None:
    data = tmp / "data.json"
    template = tmp / "custom-template.docx"
    output = tmp / "env-output.docx"
    marker = "CUSTOM ENV TEMPLATE HEADER"

    make_data(data)
    make_template(template, marker)

    env = os.environ.copy()
    env["JVC_DOCX_TEMPLATE"] = str(template)
    run_generator(data, output, env)

    assert marker in output_header_text(output), "JVC_DOCX_TEMPLATE was not used"


def assert_uses_auto_custom_template(tmp: Path) -> None:
    data = tmp / "data.json"
    output = tmp / "auto-output.docx"
    marker = "CUSTOM AUTO TEMPLATE HEADER"

    make_data(data)
    make_template(AUTO_TEMPLATE, marker)
    try:
        env = os.environ.copy()
        env.pop("JVC_DOCX_TEMPLATE", None)
        run_generator(data, output, env)
    finally:
        AUTO_TEMPLATE.unlink(missing_ok=True)

    assert marker in output_header_text(output), "templates/custom.docx was not auto-discovered"


def main() -> int:
    with tempfile.TemporaryDirectory() as tmpdir:
        tmp = Path(tmpdir)
        assert_uses_env_template(tmp)
        assert_uses_auto_custom_template(tmp)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
