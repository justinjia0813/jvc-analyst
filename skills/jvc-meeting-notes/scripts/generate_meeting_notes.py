#!/usr/bin/env python3
"""
会议纪要 Word 生成脚本
用法：python3 generate_meeting_notes.py <sections.json> --output <output.docx> [--template <template.docx>]

模板解析顺序：
1. 命令行 `--template`
2. 环境变量 `JVC_DOCX_TEMPLATE`
3. `skills/jvc-meeting-notes/templates/custom.docx`
4. 中性默认模板 `skills/jvc-meeting-notes/templates/访谈纪要模板.docx`

默认模板使用 meeting-notes 标准版式：A4、上/下 2.54cm、左/右 3.17cm、
标题居中，正文两端对齐，段前/段后 0、单倍行距，Normal 段落、run 级
Times New Roman + KaiTi 字体，并启用 doNotExpandShiftReturn，避免手动换行
短行被强行拉满。自定义模板如果提供命名样式，则优先保留用户模板样式；
如果只提供 Normal，则使用同一套直接字体格式。

sections.json 格式：
{
  "title": "2026/03/23 线上 访谈{深安锂能}",
  "sections": [
    {
      "heading": "一、公司基本情况",
      "content": "一个长句介绍..."
    },
    {
      "heading": "二、公司核心技术",
      "subsections": [
        {"heading": "技术路线", "content": "..."},
        {"heading": "核心技术", "content": "..."},
        {"heading": "达成的性能指标", "content": "..."}
      ]
    },
    ...
  ]
}
"""
import argparse
import json
import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BASE = Path(__file__).resolve().parent.parent
DEFAULT_TEMPLATE = BASE / 'templates' / '访谈纪要模板.docx'
CUSTOM_TEMPLATE = BASE / 'templates' / 'custom.docx'
DEFAULT_OUTPUT = BASE / 'output'
TEMPLATE_ENV = 'JVC_DOCX_TEMPLATE'


def resolve_template_path(explicit_template=None):
    """Resolve the template without falling back silently from bad user input."""
    if explicit_template:
        path = Path(explicit_template).expanduser()
        if not path.exists():
            raise FileNotFoundError(f'指定模板不存在: {path}')
        return path

    env_template = os.environ.get(TEMPLATE_ENV)
    if env_template:
        path = Path(env_template).expanduser()
        if not path.exists():
            raise FileNotFoundError(f'{TEMPLATE_ENV} 指向的模板不存在: {path}')
        return path

    if CUSTOM_TEMPLATE.exists():
        return CUSTOM_TEMPLATE

    if not DEFAULT_TEMPLATE.exists():
        raise FileNotFoundError(f'默认模板不存在: {DEFAULT_TEMPLATE}')
    return DEFAULT_TEMPLATE


def set_run_font(run, cn_font='KaiTi', en_font='Times New Roman', size=Pt(10), bold=False):
    """设置 run 的中英文字体"""
    run.font.size = size
    run.font.bold = bold
    run.font.name = en_font
    # 设置中文字体
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn_font)


def get_or_add_child(parent, tag):
    child = parent.find(qn(f'w:{tag}'))
    if child is None:
        child = OxmlElement(f'w:{tag}')
        parent.append(child)
    return child


def ensure_non_expanding_justify(doc):
    """Keep justified paragraphs readable when text contains manual line breaks."""
    settings = doc.settings.element
    compat = get_or_add_child(settings, 'compat')
    if compat.find(qn('w:doNotExpandShiftReturn')) is None:
        compat.append(OxmlElement('w:doNotExpandShiftReturn'))


def apply_paragraph_layout(paragraph, alignment):
    paragraph.alignment = alignment
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing = 1.0


def _style_exists(doc, style_name):
    if not style_name:
        return False
    try:
        doc.styles[style_name]
    except KeyError:
        return False
    return True


def extract_style_hints(doc):
    """Use existing template paragraphs as style hints before clearing them."""
    paragraphs = [p for p in doc.paragraphs if p.text.strip()]
    hints = {}

    preferred = {
        'title': 'JVC Note Title',
        'section': 'JVC Section Heading',
        'body': 'JVC Body',
        'subsection': 'JVC Subsection Heading',
    }
    for role, style_name in preferred.items():
        if _style_exists(doc, style_name):
            hints[role] = style_name

    positional = {
        'title': 0,
        'section': 1,
        'body': 2,
    }
    for role, index in positional.items():
        if role not in hints and len(paragraphs) > index:
            hints[role] = paragraphs[index].style.name

    if 'subsection' not in hints:
        known = {hints.get('title'), hints.get('section'), hints.get('body')}
        for paragraph in paragraphs[3:]:
            if paragraph.style.name not in known:
                hints['subsection'] = paragraph.style.name
                break
        if 'subsection' not in hints and len(paragraphs) > 5:
            hints['subsection'] = paragraphs[5].style.name
        elif 'subsection' not in hints and len(paragraphs) > 4:
            hints['subsection'] = paragraphs[4].style.name

    return {key: value for key, value in hints.items() if _style_exists(doc, value)}


def add_paragraph(doc, text, style_hints, style_key, alignment, fallback_size, fallback_bold=False):
    style_name = style_hints.get(style_key)
    if not style_name and _style_exists(doc, 'Normal'):
        style_name = 'Normal'
    use_direct_formatting = not style_name or style_name == 'Normal'
    paragraph = doc.add_paragraph() if use_direct_formatting else doc.add_paragraph(style=style_name)
    apply_paragraph_layout(paragraph, alignment)
    run = paragraph.add_run(text)
    if use_direct_formatting:
        set_run_font(run, size=fallback_size, bold=fallback_bold)
    return paragraph


def create_document(data, template_path, output_path):
    """基于模板创建会议纪要文档"""
    doc = Document(template_path)
    ensure_non_expanding_justify(doc)
    style_hints = extract_style_hints(doc)

    # 清空正文段落（保留页眉页脚和样式）
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    # 添加标题
    add_paragraph(
        doc,
        data['title'],
        style_hints,
        'title',
        WD_ALIGN_PARAGRAPH.CENTER,
        Pt(18),
        fallback_bold=True,
    )

    # 添加各板块
    for section in data['sections']:
        # 章节标题
        add_paragraph(
            doc,
            section['heading'],
            style_hints,
            'section',
            WD_ALIGN_PARAGRAPH.JUSTIFY,
            Pt(10),
            fallback_bold=True,
        )

        # 正文内容
        if 'content' in section and section['content']:
            add_paragraph(
                doc,
                section['content'],
                style_hints,
                'body',
                WD_ALIGN_PARAGRAPH.JUSTIFY,
                Pt(10),
            )

        # 子章节
        if 'subsections' in section:
            for sub in section['subsections']:
                add_paragraph(
                    doc,
                    sub['heading'],
                    style_hints,
                    'subsection',
                    WD_ALIGN_PARAGRAPH.JUSTIFY,
                    Pt(10),
                )

                if 'content' in sub and sub['content']:
                    add_paragraph(
                        doc,
                        sub['content'],
                        style_hints,
                        'body',
                        WD_ALIGN_PARAGRAPH.JUSTIFY,
                        Pt(10),
                    )

    doc.save(output_path)
    print(f'✅ 纪要已生成: {output_path}')
    print(f'模板: {template_path}')


def parse_args(argv):
    parser = argparse.ArgumentParser(
        description='Generate a meeting-note DOCX from sections JSON.'
    )
    parser.add_argument('json_path', help='sections JSON path')
    parser.add_argument('--output', default=str(DEFAULT_OUTPUT), help='output .docx path or output directory')
    parser.add_argument('--template', help='custom .docx template path')
    return parser.parse_args(argv)


def main():
    args = parse_args(sys.argv[1:])

    with open(args.json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    template = resolve_template_path(args.template)

    # 如果 output 是目录（不以 .docx 结尾），拼接 filename
    if args.output.endswith('.docx'):
        output_path = args.output
    else:
        os.makedirs(args.output, exist_ok=True)
        filename = data.get('filename', '访谈纪要.docx')
        output_path = os.path.join(args.output, filename)

    create_document(data, template, output_path)

if __name__ == '__main__':
    main()
